import re
import zipfile
from datetime import datetime
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE_DIR = Path(__file__).resolve().parent.parent
OUTPUTS_DIR = BASE_DIR / "outputs"

_ACCENT = RGBColor(0x1A, 0x52, 0x76)
_DARK = RGBColor(0x1C, 0x1C, 0x1C)
_MID = RGBColor(0x44, 0x44, 0x44)
_LIGHT = RGBColor(0x77, 0x77, 0x77)


def safe_filename(text: str) -> str:
    return re.sub(r"[^\w\s]", "", text).strip().replace(" ", "_")


def _run(paragraph, text, size_hp, color, bold=False, underline=False):
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(size_hp / 2)
    run.font.color.rgb = color
    run.bold = bold
    run.underline = underline
    return run


def _spacing(p, before_pt, after_pt):
    p.paragraph_format.space_before = Pt(before_pt)
    p.paragraph_format.space_after = Pt(after_pt)


def _bottom_border(p, color="1A5276"):
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _right_tab(p, pos_twips):
    pPr = p._p.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:pos"), str(pos_twips))
    tabs.append(tab)
    pPr.append(tabs)


def _section_heading(doc, text):
    p = doc.add_paragraph()
    _spacing(p, 12, 3)
    _bottom_border(p)
    _run(p, text.upper(), 22, _ACCENT, bold=True)
    return p


def _hyperlink(paragraph, display_text, url, size_hp=18):
    r_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hl = OxmlElement("w:hyperlink")
    hl.set(qn("r:id"), r_id)
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color_el = OxmlElement("w:color")
    color_el.set(qn("w:val"), "1A5276")
    rPr.append(color_el)
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Arial")
    rFonts.set(qn("w:hAnsi"), "Arial")
    rPr.append(rFonts)
    for tag in ("w:sz", "w:szCs"):
        el = OxmlElement(tag)
        el.set(qn("w:val"), str(size_hp))
        rPr.append(el)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    r.append(rPr)
    t = OxmlElement("w:t")
    t.text = display_text
    r.append(t)
    hl.append(r)
    paragraph._p.append(hl)


def _clear_body(doc):
    body = doc.element.body
    for p in body.findall(qn("w:p")):
        body.remove(p)


def _set_page(doc, width_tw, height_tw, margin_tw):
    sec = doc.sections[0]
    sec.page_width = Twips(width_tw)
    sec.page_height = Twips(height_tw)
    for attr in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
        setattr(sec, attr, Twips(margin_tw))


def _build_cv(tailored: dict, master_cv: dict) -> BytesIO:
    doc = Document()
    _clear_body(doc)
    _set_page(doc, 11906, 16838, 1000)

    info = master_cv["personal"]
    CONTENT_W = 9906  # 11906 - 2*1000 twips

    # Name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _spacing(p, 0, 2)
    _run(p, info.get("name", ""), 52, _DARK, bold=True)

    # Role title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _spacing(p, 0, 2)
    _run(p, tailored.get("role_title", ""), 24, _ACCENT)

    # Contact line
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _spacing(p, 0, 1)
    for text, color in [
        (info.get("email", ""), _MID),
        ("   |   ", _LIGHT),
        (info.get("phone", ""), _MID),
        ("   |   ", _LIGHT),
        (info.get("location", ""), _MID),
    ]:
        _run(p, text, 19, color)

    # Links
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _spacing(p, 0, 3)
    first = True
    for label, url in [
        ("LinkedIn", info.get("linkedin", "")),
        ("GitHub", info.get("github", "")),
        ("Portfolio", info.get("portfolio", "")),
    ]:
        if not url:
            continue
        if not first:
            _run(p, "   |   ", 18, _LIGHT)
        _hyperlink(p, label, url, 18)
        first = False

    # Summary
    _section_heading(doc, "Professional Summary")
    p = doc.add_paragraph()
    _spacing(p, 4, 4)
    _run(p, tailored.get("summary", ""), 20, _DARK)

    # Experience
    _section_heading(doc, "Experience")
    for exp in tailored.get("experience", []):
        p = doc.add_paragraph()
        _spacing(p, 6, 1)
        _right_tab(p, CONTENT_W)
        _run(p, exp.get("title", ""), 22, _DARK, bold=True)
        _run(p, "\t", 20, _DARK)
        _run(p, f"{exp.get('start', '')} – {exp.get('end', '')}", 19, _LIGHT)

        p2 = doc.add_paragraph()
        _spacing(p2, 0, 2)
        _run(p2, exp.get("company", ""), 20, _ACCENT, bold=True)
        if exp.get("type"):
            _run(p2, f"  ·  {exp['type']}", 19, _LIGHT)

        for bullet in exp.get("bullets", []):
            bp = doc.add_paragraph()
            _spacing(bp, 1, 1)
            pPr = bp._p.get_or_add_pPr()
            ind = OxmlElement("w:ind")
            ind.set(qn("w:left"), "440")
            ind.set(qn("w:hanging"), "260")
            pPr.append(ind)
            _run(bp, "•  " + bullet, 20, _DARK)

    # Projects
    if tailored.get("projects"):
        _section_heading(doc, "Projects")
        for proj in tailored["projects"]:
            p = doc.add_paragraph()
            _spacing(p, 6, 1)
            _run(p, proj.get("name", ""), 21, _DARK, bold=True)
            if proj.get("link"):
                _run(p, "  —  ", 18, _LIGHT)
                _hyperlink(p, proj["link"], proj["link"], 18)

            p2 = doc.add_paragraph()
            _spacing(p2, 0, 1)
            _run(p2, "Tech: ", 19, _MID, bold=True)
            _run(p2, ", ".join(proj.get("tech", [])), 19, _MID)

            p3 = doc.add_paragraph()
            _spacing(p3, 0, 3)
            _run(p3, proj.get("description", ""), 19, _DARK)

    # Skills
    _section_heading(doc, "Skills")
    skills = tailored.get("skills", {})
    for label, key in [
        ("Languages", "languages"),
        ("Frameworks", "frameworks"),
        ("Databases", "databases"),
        ("Tools", "tools"),
        ("Concepts", "concepts"),
    ]:
        values = skills.get(key, [])
        if values:
            p = doc.add_paragraph()
            _spacing(p, 3, 1.5)
            _run(p, f"{label}: ", 20, _DARK, bold=True)
            _run(p, ", ".join(values), 20, _MID)

    # Education
    _section_heading(doc, "Education")
    for edu in master_cv.get("education", []):
        p = doc.add_paragraph()
        _spacing(p, 5, 1)
        _right_tab(p, CONTENT_W)
        _run(p, f"{edu.get('degree', '')} — {edu.get('field', '')}", 21, _DARK, bold=True)
        _run(p, "\t", 20, _DARK)
        _run(p, edu.get("year", ""), 19, _LIGHT)

        p2 = doc.add_paragraph()
        _spacing(p2, 0, 3)
        _run(p2, edu.get("institution", ""), 20, _ACCENT)
        _run(p2, f"  ·  {edu.get('status', '')}", 19, _LIGHT)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def _build_cover_letter(tailored: dict, master_cv: dict) -> BytesIO:
    doc = Document()
    _clear_body(doc)
    _set_page(doc, 11906, 16838, 1440)

    info = master_cv["personal"]

    p = doc.add_paragraph()
    _spacing(p, 0, 1)
    _run(p, info.get("name", ""), 26, _DARK, bold=True)

    p = doc.add_paragraph()
    _spacing(p, 0, 1)
    for text, color in [
        (info.get("email", ""), _MID),
        ("   |   ", _LIGHT),
        (info.get("phone", ""), _MID),
        ("   |   ", _LIGHT),
        (info.get("location", ""), _MID),
    ]:
        _run(p, text, 20, color)

    p = doc.add_paragraph()
    _spacing(p, 0, 3)
    _bottom_border(p)

    p = doc.add_paragraph()
    _spacing(p, 15, 4)
    _run(p, datetime.today().strftime("%d %B %Y"), 20, _LIGHT)

    p = doc.add_paragraph()
    _spacing(p, 4, 3)
    _run(p, "Dear Hiring Manager,", 21, _DARK)

    for para in tailored.get("cover_letter", "").strip().split("\n\n"):
        if para.strip():
            p = doc.add_paragraph()
            _spacing(p, 0, 9)
            _run(p, para.strip(), 21, _DARK)

    p = doc.add_paragraph()
    _spacing(p, 6, 1)
    _run(p, "Regards,", 21, _DARK)

    p = doc.add_paragraph()
    _spacing(p, 0, 0)
    _run(p, info.get("name", ""), 21, _DARK, bold=True)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def generate_documents(tailored: dict, master_cv: dict) -> Path:
    OUTPUTS_DIR.mkdir(exist_ok=True)

    role = safe_filename(tailored.get("role_title", "Role"))
    company = safe_filename(tailored.get("company", "Company"))

    zip_out = OUTPUTS_DIR / f"Application_{role}_{company}.zip"

    cv_buf = _build_cv(tailored, master_cv)
    cl_buf = _build_cover_letter(tailored, master_cv)

    with zipfile.ZipFile(zip_out, "w") as zf:
        zf.writestr(f"CV_{role}_{company}.docx", cv_buf.getvalue())
        zf.writestr(f"CoverLetter_{role}_{company}.docx", cl_buf.getvalue())

    return zip_out
